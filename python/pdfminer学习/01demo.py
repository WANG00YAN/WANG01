import sys
import os
sys.path.append("/media/wang/103E0A4F103E0A4F1/code/learn/python/")
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams,LTTextBoxHorizontal,LTImage,LTCurve,LTFigure
from pdfminer.pdfpage import PDFTextExtractionNotAllowed,PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from docx import Document
from common.color import PrintColor




document = Document()
'''
pip install pdfminer3k
pip install pdfminer.six 安装这个引入的内容不会报错
'''

class CPdf2TxtManager():

    def changePdfToText(self, filePath):
        # 以二进制读模式打开
        file = open(path, 'rb')
        #用文件对象来创建一个pdf文档分析器
        praser = PDFParser(file)
        # 创建一个PDF文档对象存储文档结构,提供密码初始化，没有就不用传该参数
        doc = PDFDocument(praser, password='')
        ##检查文件是否允许文本提取
        if not doc.is_extractable:
            raise PDFTextExtractionNotAllowed

        # 创建PDf 资源管理器 来管理共享资源，#caching = False不缓存
        rsrcmgr = PDFResourceManager(caching = False)
        # 创建一个PDF设备对象
        laparams = LAParams()
        # 创建一个PDF页面聚合对象
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解析器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        # 获得文档的目录（纲要）,文档没有纲要会报错
        #PDF文档没有目录时会报：raise PDFNoOutlines  pdfminer.pdfdocument.PDFNoOutlines
        # print(doc.get_outlines())

        # 获取page列表
        print(PDFPage.get_pages(doc))
        # 用来计数页面，图片，曲线，figure，水平文本框等对象的数量
        num_page, num_image, num_curve, num_figure, num_TextBoxHorizontal = 0, 0, 0, 0, 0
        # 循环遍历列表，每次处理一个page的内容
        for page in PDFPage.create_pages(doc):
            num_page += 1  # 页面增一
            # 利用解释器的process_page()方法解析读取单独页数
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            fileNames = os.path.splitext(filePath)
            # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象
            # 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等
            layout = sorted(layout,key=lambda x:-x.y0)
            # for x in layout:
            #     print(f"({x.x0},{x.y0})  ({x.x1},{x.y1})")
            # self.analyse_layout(layout)
          
            self.analyse_layout(layout)
            break
            
            # for x in layout:
            #     if hasattr(x, "get_text") or isinstance(x, LTTextBoxHorizontal):
            #         with open(fileNames[0] + '.txt','a+') as f:
            #             # 将' '替换成u' '空格，这个 就是&nbps空格
            #             results = x.get_text().replace(u' ', u' ')
            #             f.write(results + '\n')
            #         document.add_paragraph(
            #             results, style='ListBullet'  # 添加段落，样式为unordered list类型
            #         )
            #     document.save('./demo1.docx')  # 保存这个文档

            #     # 如果x是水平文本对象的话
            #     if isinstance(x, LTTextBoxHorizontal):
            #         num_TextBoxHorizontal += 1  # 水平文本框对象增一
            #     if isinstance(x, LTImage):  # 图片对象
            #         num_image += 1
            #     if isinstance(x, LTCurve):  # 曲线对象
            #         num_curve += 1
            #     if isinstance(x, LTFigure):  # figure对象
            #         num_figure += 1

        print('对象数量：%s,页面数：%s,图片数：%s,曲线数：%s,'
              '水平文本框：%s,'%(num_figure,num_page,num_image,num_curve,num_TextBoxHorizontal))


    def analyse_layout(self, layout):
        
        for x in layout:
            print(type(x))
            if isinstance(x, LTTextBoxHorizontal):
                with open('test.txt', 'a') as f:
                    print(PrintColor.yellow)
                    print(x.get_text().strip())
                    print(PrintColor.red)
                    print("="*100)
                    print(PrintColor.end)
                    f.write(x.get_text().strip())
            



    
if __name__ == '__main__':
    path = '/media/wang/103E0A4F103E0A4F1/zhiwang/TKDE18_gAnswer.pdf'
    pdf2TxtManager = CPdf2TxtManager()
    pdf2TxtManager.changePdfToText(path)
    print('ok,解析pdf结束!')